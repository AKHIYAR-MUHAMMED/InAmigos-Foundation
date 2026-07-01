import os
import sys
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from fpdf import FPDF

# Custom PDF Class mirroring NGOReportPDF styles
class NGOGrowthPDF(FPDF):
    def header(self):
        # Top margin spacing / header
        pass
        
    def footer(self):
        # Footer layout with page numbers
        self.set_y(-15)
        self.set_font("helvetica", "I", 8)
        self.set_text_color(74, 85, 104) # Slate Color
        text = f"InAmigos Foundation | AI Insights for NGO Growth (2026)  -  Page {self.page_no()} of {{nb}}"
        self.cell(0, 10, text, border=0, align="R")

def generate_docx():
    doc = Document()
    
    # 1-inch margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Color palette matching InAmigos NGO theme
    NAVY = RGBColor(26, 54, 93)      # #1A365D - Main headings
    TEAL = RGBColor(49, 151, 149)    # #319795 - Accent
    CHARCOAL = RGBColor(45, 55, 72)  # #2D3748 - Body Text
    SLATE = RGBColor(74, 85, 104)    # #4A5568 - Secondary Text
    
    # Helper to add a heading with custom style, spacing, and teal bottom line
    def add_custom_heading(text, level, space_before, space_after):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.keep_with_next = True
        
        run = p.add_run(text)
        run.bold = True
        if level == 1:
            run.font.size = Pt(15)
            run.font.color.rgb = NAVY
            run.font.name = 'Calibri'
            # Bottom accent line XML element
            pPr = p._p.get_or_add_pPr()
            pbdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '12') # 1.5 pt
            bottom.set(qn('w:space'), '4')
            bottom.set(qn('w:color'), '319795') # Teal accent
            pbdr.append(bottom)
            pPr.append(pbdr)
        elif level == 2:
            run.font.size = Pt(12.5)
            run.font.color.rgb = TEAL
            run.font.name = 'Calibri'
        return p

    def add_body_paragraph(text, space_after=6, bold_prefix=None):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = 1.15
        
        if bold_prefix:
            run_prefix = p.add_run(bold_prefix)
            run_prefix.bold = True
            run_prefix.font.size = Pt(11)
            run_prefix.font.color.rgb = CHARCOAL
            run_prefix.font.name = 'Calibri'
            
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.color.rgb = CHARCOAL
        run.font.name = 'Calibri'
        return p

    def add_bullet_point(text, bold_prefix=None):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        
        if bold_prefix:
            run_prefix = p.add_run(bold_prefix)
            run_prefix.bold = True
            run_prefix.font.size = Pt(11)
            run_prefix.font.color.rgb = CHARCOAL
            run_prefix.font.name = 'Calibri'
            
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.color.rgb = CHARCOAL
        run.font.name = 'Calibri'
        return p

    # Header / Title Block
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(12)
    p_title.paragraph_format.space_after = Pt(4)
    run_title = p_title.add_run("AI INSIGHTS FOR NGO GROWTH")
    run_title.bold = True
    run_title.font.size = Pt(22)
    run_title.font.color.rgb = NAVY
    run_title.font.name = 'Calibri'
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(18)
    run_sub = p_sub.add_run("5 AI-Powered Strategies to Enhance Outreach, Engagement, Fundraising, and Brand Presence")
    run_sub.font.size = Pt(11)
    run_sub.font.italic = True
    run_sub.font.color.rgb = SLATE
    run_sub.font.name = 'Calibri'

    # Line Divider
    p_div = doc.add_paragraph()
    p_div.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_div.paragraph_format.space_after = Pt(14)
    run_div = p_div.add_run("—" * 50)
    run_div.font.color.rgb = TEAL

    # Executive Summary (Styled Callout Box)
    add_custom_heading("Executive Summary", 1, 10, 6)
    
    summary_table = doc.add_table(rows=1, cols=1)
    summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = summary_table.cell(0, 0)
    cell.width = Inches(6.5)
    
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F7FAFC"/>')
    tcPr.append(shd)
    
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:left w:val="single" w:sz="24" w:space="0" w:color="319795"/><w:top w:val="none"/><w:right w:val="none"/><w:bottom w:val="none"/></w:tcBorders>')
    tcPr.append(tcBorders)
    
    p_sum = cell.paragraphs[0]
    p_sum.paragraph_format.left_indent = Inches(0.15)
    p_sum.paragraph_format.right_indent = Inches(0.15)
    p_sum.paragraph_format.space_before = Pt(8)
    p_sum.paragraph_format.space_after = Pt(8)
    p_sum.paragraph_format.line_spacing = 1.15
    run_sum = p_sum.add_run(
        "For modern non-profit organizations like the InAmigos Foundation, social impact relies heavily on the "
        "efficiency of outreach and the sustainability of donor funding. In 2026, Artificial Intelligence (AI) has shifted "
        "from a futuristic tool to an essential operational asset. Rather than replacing the essential empathy of NGO work, "
        "AI enables small teams to amplify their reach, build deeper relationships with volunteers, automate content creation, "
        "and maximize fundraising retention. This report details 5 concrete, actionable strategies that the InAmigos Foundation "
        "can implement to drive organic community growth and secure sustainable funding."
    )
    run_sum.font.size = Pt(10)
    run_sum.font.color.rgb = CHARCOAL
    run_sum.font.name = 'Calibri'
    run_sum.font.italic = True

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # 1. Hyper-Personalized Donor Communication
    add_custom_heading("1. AI-Driven Hyper-Personalized Donor Communication", 1, 14, 6)
    add_body_paragraph(
        "Traditional fundraising outreach often relies on generic mass email campaigns. However, modern donors demand "
        "relevance and transparency regarding the projects they support. AI segmentation makes hyper-personalization scalable:"
    )
    add_bullet_point(
        "AI algorithms analyze historical donor donation records, specific project affinity (e.g., child education "
        "vs. eco tree plantation drives), and email click patterns to group donors automatically.",
        bold_prefix="Behavioral Segmentation: "
    )
    add_bullet_point(
        "Instead of broad monthly newsletters, generative models draft custom appeal updates. For instance, a donor who "
        "supported education receives an update highlighting specific village classrooms, while an environmental advocate "
        "receives reports on carbon offsets and forest sapling growth.",
        bold_prefix="Tailored Content Copywriting: "
    )
    add_bullet_point(
        "Utilizing tools like HubSpot AI, Mailchimp AI, or Claude API to draft email templates, which are then reviewed "
        "by InAmigos staff to ensure authentic human tone.",
        bold_prefix="Actionable Deployment: "
    )

    # 2. Conversational AI Chatbots for Volunteer Support
    add_custom_heading("2. Intelligent Conversational Chatbots for Volunteer Onboarding", 1, 14, 6)
    add_body_paragraph(
        "Volunteer inquiries represent a significant administrative bottleneck. A delayed response can lead a passionate "
        "volunteer to lose interest. Conversational AI can bridge this response gap immediately:"
    )
    add_bullet_point(
        "An AI-powered assistant (integrated on the InAmigos website or WhatsApp Business) interacts with potential "
        "volunteers 24/7, answering common FAQs regarding locations, commitment hours, and program requirements.",
        bold_prefix="Instant Response & FAQs: "
    )
    add_bullet_point(
        "The bot leads applicants through a simple screening workflow—asking for their background, skills (e.g., teaching, "
        "graphic design, field coordination), and availability, and automatically updates the NGO's coordination sheet.",
        bold_prefix="Automated Screening & Intake: "
    )
    add_bullet_point(
        "Setting up Dialogflow, Landbot, or a custom WhatsApp integration linked to an OpenAI assistant. This ensures "
        "prospective helpers are immediately engaged within seconds of expressing interest.",
        bold_prefix="Actionable Deployment: "
    )

    # Page Break for layout structure
    doc.add_page_break()

    # 3. AI-Driven Social Media Content & Reel Automation
    add_custom_heading("3. AI-Powered Social Media Content & Brand Presence", 1, 12, 6)
    add_body_paragraph(
        "A strong, active online presence is crucial for establishing credibility. Generative AI tools allow small communication "
        "teams to maintain a high-quality, high-frequency posting schedule across LinkedIn, Instagram, and YouTube:"
    )
    add_bullet_point(
        "AI writing assistants analyze raw notes, field photographs, and speech outlines from InAmigos NGO projects to "
        "generate polished LinkedIn posts, caption copy, and educational summaries.",
        bold_prefix="Multi-Format Copy Generation: "
    )
    add_bullet_point(
        "Using generative video scripts and outlines to script engaging Instagram Reels or YouTube Shorts. Generative tools "
        "can automatically add dynamic captions, match background royalty-free audio, and highlight engaging key moments.",
        bold_prefix="Micro-Video Scripting: "
    )
    add_bullet_point(
        "Leveraging Canva Magic Write, ChatGPT, and CapCut AI for graphic layouts and script brainstorming; utilizing Buffer AI "
        "to schedule posts at peak engagement times.",
        bold_prefix="Actionable Deployment: "
    )

    # 4. Predictive Churn Modeling for Donor Retention
    add_custom_heading("4. Predictive Churn Modeling & Retention Campaigns", 1, 14, 6)
    add_body_paragraph(
        "Acquiring new donors is far more expensive than retaining existing ones. Predictive machine learning allows NGOs "
        "to identify 'at-risk' donors before they cancel their subscriptions or cease giving:"
    )
    add_bullet_point(
        "By analyzing simple variables like newsletter open rates, skipped monthly contributions, or a lack of attendance "
        "at recent events, predictive algorithms flag donors exhibiting warning behaviors.",
        bold_prefix="Risk Profiling: "
    )
    add_bullet_point(
        "Once a donor is flagged with a high 'churn score,' the system alerts the fundraising team. Staff can intervene "
        "specifically with a personalized phone call, a hand-written thank-you card, or a customized impact report.",
        bold_prefix="Targeted Human Intervention: "
    )
    add_bullet_point(
        "Employing Dataro Predict or building a lightweight Python script using pandas and scikit-learn to analyze "
        "existing donor Excel sheets or CRM databases.",
        bold_prefix="Actionable Deployment: "
    )

    # 5. AI-Assisted CSR & Grant Proposal Writing
    add_custom_heading("5. AI-Assisted Corporate Sponsorship & CSR Proposal Drafting", 1, 14, 6)
    add_body_paragraph(
        "NGO growth is often limited by the time spent seeking grants. Corporate Social Responsibility (CSR) proposals "
        "require rigorous formatting, statistical evidence, and alignment with corporate guidelines:"
    )
    add_bullet_point(
        "Generative AI models serve as highly-capable drafting assistants. By loading the model with historical successful proposals, "
        "annual reports, and local impact metrics, the AI can draft structured, professional proposals tailored to corporate sponsors.",
        bold_prefix="Intelligent Drafting Support: "
    )
    add_bullet_point(
        "AI can quickly format and adjust proposals to match different structures (e.g., modifying a 10-page detailed grant request "
        "into a 1-page corporate pitch deck summary), ensuring compliance with varying corporate guidelines.",
        bold_prefix="Formatting & Adaptation: "
    )
    add_bullet_point(
        "Developing a custom Claude or ChatGPT Project workspace loaded with past proposal PDF templates, allowing staff "
        "to brainstorm pitches and outline drafts in minutes.",
        bold_prefix="Actionable Deployment: "
    )

    # Comparative Metrics Table
    add_custom_heading("Strategic Implementation & Metrics Comparison", 1, 18, 6)
    add_body_paragraph(
        "To help InAmigos Foundation plan its implementation, this comparison table outlines the primary focus, "
        "primary tools, required effort, and expected metrics for each of the 5 strategies:"
    )

    # Create table
    table = doc.add_table(rows=6, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    # Table Header
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Strategy"
    hdr_cells[1].text = "Focus Area"
    hdr_cells[2].text = "Recommended AI Tools"
    hdr_cells[3].text = "Difficulty & Impact"
    
    col_widths = [Inches(1.8), Inches(1.3), Inches(1.8), Inches(1.6)]
    for row in table.rows:
        for idx, width in enumerate(col_widths):
            row.cells[idx].width = width
            
    # Header styling (Navy background, white text)
    for cell in hdr_cells:
        tcPr = cell._tc.get_or_add_tcPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1A365D"/>')
        tcPr.append(shd)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(9.5)
                run.font.name = 'Calibri'
                
    # Rows data
    row_data = [
        ("Hyper-Personalized Communication", "Fundraising & Outreach", "HubSpot AI, Mailchimp, Claude API", "Low Effort | High Impact"),
        ("Conversational AI Chatbots", "Volunteer Engagement", "Dialogflow, Landbot, OpenAI API", "Medium Effort | High Impact"),
        ("Social Media Content Automation", "Online Brand Presence", "Canva Magic Write, CapCut, Buffer", "Low Effort | Medium Impact"),
        ("Predictive Donor Churn Modeling", "Fundraising & Retention", "Dataro Predict, Python scripts", "Medium Effort | High Impact"),
        ("AI CSR Proposal Drafting", "Grant Acquisitions", "Claude Projects, Custom GPTs", "Low Effort | High Impact")
    ]
    
    for idx, data in enumerate(row_data):
        row = table.rows[idx + 1]
        row.cells[0].text = data[0]
        row.cells[1].text = data[1]
        row.cells[2].text = data[2]
        row.cells[3].text = data[3]
        
        bg_color = "F7FAFC" if idx % 2 == 0 else "FFFFFF"
        
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg_color}"/>')
            tcPr.append(shd)
            
            tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:bottom w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/></w:tcBorders>')
            tcPr.append(tcBorders)
            
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)
                for run in p.runs:
                    run.font.size = Pt(9)
                    run.font.color.rgb = CHARCOAL
                    run.font.name = 'Calibri'

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Next Steps Conclusion
    add_custom_heading("Conclusion & Immediate Next Steps", 1, 12, 6)
    add_body_paragraph(
        "Implementing AI in InAmigos Foundation's operations does not require massive capital investments. The foundation can "
        "adopt a phased transition starting with the lowest-effort, highest-impact strategies:"
    )
    add_bullet_point(
        "Establish an AI copywriting sandbox utilizing Claude Projects or ChatGPT, loaded with InAmigos brochures, "
        "history, and successful templates, to draft volunteer newsletters and LinkedIn posts immediately.",
        bold_prefix="Phase 1 (Week 1): "
    )
    add_bullet_point(
        "Design and launch a simple, rules-based chatbot via Landbot or WhatsApp Business to answer basic volunteer registration "
        "FAQs, gradually upgrading it to a conversational AI model as database query volumes grow.",
        bold_prefix="Phase 2 (Month 1): "
    )
    add_bullet_point(
        "Perform a quality audit of the active CRM/donor Excel database. Ensure fields are clean and uniform (e.g., separate columns "
        "for donation dates, project category, and donation size) to prepare for future predictive donor models.",
        bold_prefix="Phase 3 (Month 2): "
    )

    # Footer Setup
    for section in doc.sections:
        footer = section.footer
        p_foot = footer.paragraphs[0]
        p_foot.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_foot.paragraph_format.space_before = Pt(12)
        run_foot = p_foot.add_run("InAmigos Foundation | AI Insights for NGO Growth (2026)  —  Page ")
        run_foot.font.size = Pt(8.5)
        run_foot.font.color.rgb = SLATE
        run_foot.font.name = 'Calibri'
        
        # XML to generate current page number field in Word
        fldChar1 = parse_xml(r'<w:fldChar %s w:fldCharType="begin"/>' % nsdecls('w'))
        instrText = parse_xml(r'<w:instrText %s xml:space="preserve"> PAGE </w:instrText>' % nsdecls('w'))
        fldChar2 = parse_xml(r'<w:fldChar %s w:fldCharType="separate"/>' % nsdecls('w'))
        fldChar3 = parse_xml(r'<w:fldChar %s w:fldCharType="end"/>' % nsdecls('w'))
        p_foot._p.append(fldChar1)
        p_foot._p.append(instrText)
        p_foot._p.append(fldChar2)
        p_foot._p.append(fldChar3)

    filename = "AI_NGO_Growth_Insights.docx"
    doc.save(filename)
    print(f"DOCX successfully generated: {os.path.abspath(filename)}")

def generate_pdf():
    pdf = NGOGrowthPDF(orientation="P", unit="mm", format="A4")
    pdf.set_margins(25.4, 25.4, 25.4) # 1 inch margins
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.alias_nb_pages()
    
    # Colors
    NAVY = (26, 54, 93)
    TEAL = (49, 151, 149)
    CHARCOAL = (45, 55, 72)
    SLATE = (74, 85, 104)
    LIGHT_GRAY = (247, 250, 252)
    BORDER_GRAY = (226, 232, 240)
    
    pdf.add_page()
    
    # Document Title Block
    pdf.set_font("helvetica", "B", 18)
    pdf.set_text_color(*NAVY)
    pdf.cell(0, 10, "AI INSIGHTS FOR NGO GROWTH", new_x="LMARGIN", new_y="NEXT", align="C")
    
    pdf.set_font("helvetica", "I", 10.5)
    pdf.set_text_color(*SLATE)
    pdf.cell(0, 6, "5 AI-Powered Strategies for InAmigos Foundation", new_x="LMARGIN", new_y="NEXT", align="C")
    
    pdf.ln(3)
    pdf.set_draw_color(*TEAL)
    pdf.set_line_width(0.5)
    pdf.line(25.4, pdf.get_y(), 210 - 25.4, pdf.get_y())
    pdf.ln(5)
    
    # Helper functions
    def add_h1(text):
        pdf.set_font("helvetica", "B", 12.5)
        pdf.set_text_color(*NAVY)
        pdf.ln(4)
        pdf.cell(0, 7, text, new_x="LMARGIN", new_y="NEXT")
        pdf.set_draw_color(*TEAL)
        pdf.set_line_width(0.3)
        pdf.line(25.4, pdf.get_y(), 210 - 25.4, pdf.get_y())
        pdf.ln(2.5)
        
    def add_body(text, bold_prefix=None):
        pdf.set_font("helvetica", "", 9.5)
        pdf.set_text_color(*CHARCOAL)
        if bold_prefix:
            pdf.set_font("helvetica", "B", 9.5)
            pdf.write(4.5, bold_prefix)
            pdf.set_font("helvetica", "", 9.5)
        pdf.write(4.5, text)
        pdf.ln(5.5)
        
    def add_bullet(text, bold_prefix=None):
        pdf.set_font("helvetica", "", 9.5)
        pdf.set_text_color(*CHARCOAL)
        x = pdf.get_x()
        y = pdf.get_y()
        # Draw a tiny teal square bullet
        pdf.set_fill_color(*TEAL)
        pdf.set_draw_color(*TEAL)
        pdf.rect(x + 2, y + 1.5, 1.2, 1.2, style="F")
        
        pdf.set_x(x + 5.5)
        if bold_prefix:
            pdf.set_font("helvetica", "B", 9.5)
            pdf.write(4.5, bold_prefix)
            pdf.set_font("helvetica", "", 9.5)
        pdf.write(4.5, text)
        pdf.ln(5.5)

    # Executive Summary H1
    add_h1("Executive Summary")
    
    summary_text = (
        "For modern non-profit organizations like the InAmigos Foundation, social impact relies heavily on the "
        "efficiency of outreach and the sustainability of donor funding. In 2026, Artificial Intelligence (AI) has shifted "
        "from a futuristic tool to an essential operational asset. Rather than replacing the essential empathy of NGO work, "
        "AI enables small teams to amplify their reach, build deeper relationships with volunteers, automate content creation, "
        "and maximize fundraising retention. This report details 5 concrete, actionable strategies that the InAmigos Foundation "
        "can implement to drive organic community growth and secure sustainable funding."
    )
    
    box_x = 25.4
    box_y = pdf.get_y()
    
    pdf.set_fill_color(*LIGHT_GRAY)
    pdf.set_draw_color(*LIGHT_GRAY)
    pdf.set_font("helvetica", "I", 9)
    pdf.set_text_color(*CHARCOAL)
    
    # Print callout box text
    pdf.multi_cell(159.2, 4.5, summary_text, border=0, align="L", fill=True)
    box_end_y = pdf.get_y()
    
    # Draw left teal border line
    pdf.set_draw_color(*TEAL)
    pdf.set_line_width(1.0)
    pdf.line(box_x, box_y, box_x, box_end_y)
    pdf.ln(3)
    
    # Strategy 1
    add_h1("1. AI-Driven Hyper-Personalized Donor Communication")
    add_body("Traditional fundraising campaigns often suffer from low response rates due to generic appeal emails. AI makes personalization scalable:")
    add_bullet("AI models analyze donor histories, transaction size, and project affinities (e.g., education vs. tree planting) to cluster donor bases.", bold_prefix="Smart Segmentation: ")
    add_bullet("Generative tools draft custom newsletters showing impact metrics specifically for the projects each donor cares about.", bold_prefix="Impact Copywriting: ")
    add_bullet("Using HubSpot, Mailchimp AI, or Claude API to generate personalized drafts for staff to review.", bold_prefix="Actionable Tool: ")
    
    # Strategy 2
    add_h1("2. Intelligent Conversational Chatbots for Volunteer Onboarding")
    add_body("Coordinating volunteer signups manually creates huge administrative delays. Conversational AI bridges the gap immediately:")
    add_bullet("An AI chatbot on the web or WhatsApp acts as a 24/7 assistant, answering FAQs and guiding registrations.", bold_prefix="24/7 Desk: ")
    add_bullet("The chatbot collects volunteer availability, interests, and skills (e.g., teaching, graphic design) automatically.", bold_prefix="Automated Screening: ")
    add_bullet("Setting up Landbot, Dialogflow, or WhatsApp API with OpenAI assistants to capture prospective leads instantly.", bold_prefix="Actionable Tool: ")
    
    # Page Break
    pdf.add_page()
    
    # Strategy 3
    add_h1("3. AI-Powered Social Media Content & Brand Presence")
    add_body("Social credibility requires consistent digital posting. Generative AI helps small outreach teams work at scale:")
    add_bullet("AI tools process field photos and coordinator log transcripts to outline clean LinkedIn posts and caption copy.", bold_prefix="Drafting Assistant: ")
    add_bullet("Generative outlines and templates help script reels and shorts, automatically synching subtitles and audio tags.", bold_prefix="Reels Scripting: ")
    add_bullet("Employing Canva Magic Write, ChatGPT, and CapCut; scheduling with Buffer/Hootsuite AI to optimize engagement.", bold_prefix="Actionable Tool: ")
    
    # Strategy 4
    add_h1("4. Predictive Churn Modeling & Retention Campaigns")
    add_body("Retaining existing regular donors is more cost-efficient than marketing to secure new ones:")
    add_bullet("By identifying donor patterns (e.g., skipped monthly gifts, dropped email opens), predictive systems flag at-risk records.", bold_prefix="Risk Profiling: ")
    add_bullet("Flagged records trigger a personalized call or handwritten card from InAmigos staff to check in and show appreciation.", bold_prefix="Empathetic Outreach: ")
    add_bullet("Using Dataro Predict, or developing basic Python spreadsheets analysis scripts using pandas.", bold_prefix="Actionable Tool: ")
    
    # Strategy 5
    add_h1("5. AI-Assisted Corporate Sponsorship & CSR Proposal Drafting")
    add_body("CSR and government grants are highly detailed and require specific language patterns:")
    add_bullet("Generative AI models trained on past proposals draft customized grant proposals that align with corporate values.", bold_prefix="Proposal Drafting: ")
    add_bullet("AI helps rewrite, summarize, or reformulate complex proposals to fit multiple corporate application templates.", bold_prefix="Formatting Utility: ")
    add_bullet("Setting up ChatGPT Custom GPTs or Claude Projects loaded with previous proposals to draft templates rapidly.", bold_prefix="Actionable Tool: ")
    
    # Table Page or Table Section
    add_h1("Strategic Implementation & Metrics Comparison")
    pdf.ln(1)
    
    # Table widths
    col_widths = [45, 32, 45, 37.2]
    row_height = 7
    
    # Header
    pdf.set_fill_color(*NAVY)
    pdf.set_text_color(255, 255, 255)
    pdf.set_font("helvetica", "B", 8)
    pdf.cell(col_widths[0], row_height, "Strategy", border=1, align="L", fill=True)
    pdf.cell(col_widths[1], row_height, "Focus Area", border=1, align="L", fill=True)
    pdf.cell(col_widths[2], row_height, "Recommended Tools", border=1, align="L", fill=True)
    pdf.cell(col_widths[3], row_height, "Difficulty & Impact", border=1, align="L", fill=True)
    pdf.ln()
    
    # Rows
    rows_data = [
        ("Hyper-Personalized Comm.", "Fundraising / Outreach", "HubSpot, Mailchimp, Claude API", "Low Effort | High Impact"),
        ("Conversational Chatbots", "Volunteer Engagement", "Dialogflow, Landbot, OpenAI", "Medium Effort | High Impact"),
        ("Social Media Automation", "Online Brand Presence", "Canva Magic Write, CapCut", "Low Effort | Medium Impact"),
        ("Predictive Churn Modeling", "Fundraising & Retention", "Dataro Predict, Pandas", "Medium Effort | High Impact"),
        ("AI Proposal Drafting", "Grant Acquisitions", "Claude Projects, Custom GPTs", "Low Effort | High Impact")
    ]
    
    pdf.set_text_color(*CHARCOAL)
    pdf.set_font("helvetica", "", 7.5)
    
    for idx, (strategy, focus, tools, diff) in enumerate(rows_data):
        bg = LIGHT_GRAY if idx % 2 == 0 else (255, 255, 255)
        pdf.set_fill_color(*bg)
        
        # Multi-cell auto-alignment
        x = pdf.get_x()
        y = pdf.get_y()
        
        pdf.multi_cell(col_widths[0], 5, strategy, border=1, align="L", fill=True)
        h1 = pdf.get_y() - y
        
        pdf.set_xy(x + col_widths[0], y)
        pdf.multi_cell(col_widths[1], 5, focus, border=1, align="L", fill=True)
        h2 = pdf.get_y() - y
        
        pdf.set_xy(x + col_widths[0] + col_widths[1], y)
        pdf.multi_cell(col_widths[2], 5, tools, border=1, align="L", fill=True)
        h3 = pdf.get_y() - y
        
        pdf.set_xy(x + col_widths[0] + col_widths[1] + col_widths[2], y)
        pdf.multi_cell(col_widths[3], 5, diff, border=1, align="L", fill=True)
        h4 = pdf.get_y() - y
        
        max_h = max(h1, h2, h3, h4)
        pdf.set_xy(25.4, y + max_h)
        
    pdf.ln(3)
    
    # Conclusion
    pdf.set_font("helvetica", "B", 10.5)
    pdf.set_text_color(*TEAL)
    pdf.cell(0, 5, "Conclusion & Immediate Next Steps", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(1)
    
    add_bullet("Establish an AI copywriting sandbox utilizing Claude Projects or ChatGPT, loaded with InAmigos brochures, history, and successful templates, to draft newsletters and social posts.", bold_prefix="Phase 1: ")
    add_bullet("Launch a simple, rules-based chatbot via Landbot or WhatsApp to handle volunteer registration FAQs.", bold_prefix="Phase 2: ")
    add_bullet("Perform a clean audit of the active donor Excel database to prepare fields for predictive metrics.", bold_prefix="Phase 3: ")

    filename = "AI_NGO_Growth_Insights.pdf"
    pdf.output(filename)
    print(f"PDF successfully generated: {os.path.abspath(filename)}")

if __name__ == "__main__":
    generate_docx()
    generate_pdf()
