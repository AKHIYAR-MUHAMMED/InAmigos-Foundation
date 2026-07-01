import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def create_report():
    doc = Document()
    
    # Page setup - 1 inch margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Define color palette
    NAVY = RGBColor(26, 54, 93)      # #1A365D - Main headings
    TEAL = RGBColor(49, 151, 149)    # #319795 - Subheadings / Accent
    CHARCOAL = RGBColor(45, 55, 72)  # #2D3748 - Body Text
    SLATE = RGBColor(74, 85, 104)    # #4A5568 - Secondary Text
    
    # Helper to add a heading with custom color, size, and spacing
    def add_custom_heading(text, level, space_before, space_after):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.keep_with_next = True
        
        run = p.add_run(text)
        run.bold = True
        if level == 1:
            run.font.size = Pt(16)
            run.font.color.rgb = NAVY
            run.font.name = 'Calibri'
            # Add bottom border/line XML element for style
            pPr = p._p.get_or_add_pPr()
            pbdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '12') # 1.5 pt
            bottom.set(qn('w:space'), '4')
            bottom.set(qn('w:color'), '319795') # Teal accent
            pbdr.append(bottom)
            pPr.append(pbdr)
        elif level == 2:
            run.font.size = Pt(13)
            run.font.color.rgb = TEAL
            run.font.name = 'Calibri'
        return p

    def add_body_paragraph(text, space_after=6, bold_prefix=None):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = 1.15
        
        if bold_prefix:
            run_prefix = p.add_run(bold_prefix)
            run_prefix.bold = True
            run_prefix.font.size = Pt(11)
            run_prefix.font.color.rgb = CHARCOAL
            run_prefix.font.name = 'Calibri'
            
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.color.rgb = CHARCOAL
        run.font.name = 'Calibri'
        return p

    def add_bullet_point(text, bold_prefix=None):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        
        if bold_prefix:
            run_prefix = p.add_run(bold_prefix)
            run_prefix.bold = True
            run_prefix.font.size = Pt(11)
            run_prefix.font.color.rgb = CHARCOAL
            run_prefix.font.name = 'Calibri'
            
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.color.rgb = CHARCOAL
        run.font.name = 'Calibri'
        return p

    # Title Page/Header Section
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(12)
    p_title.paragraph_format.space_after = Pt(4)
    run_title = p_title.add_run("ARTIFICIAL INTELLIGENCE IN NGO OPERATIONS")
    run_title.bold = True
    run_title.font.size = Pt(22)
    run_title.font.color.rgb = NAVY
    run_title.font.name = 'Calibri'
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(18)
    run_sub = p_sub.add_run("Enhancing Fundraising, Donor Retention, and Operational Efficiency in 2026")
    run_sub.font.size = Pt(12)
    run_sub.font.italic = True
    run_sub.font.color.rgb = SLATE
    run_sub.font.name = 'Calibri'

    # Divider line
    p_div = doc.add_paragraph()
    p_div.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_div.paragraph_format.space_after = Pt(18)
    run_div = p_div.add_run("—" * 50)
    run_div.font.color.rgb = TEAL

    # Executive Summary (Styled callout box)
    add_custom_heading("Executive Summary", 1, 12, 6)
    
    summary_table = doc.add_table(rows=1, cols=1)
    summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = summary_table.cell(0, 0)
    cell.width = Inches(6.5)
    
    # Style callout box background and borders
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F7FAFC"/>') # Light off-white background
    tcPr.append(shd)
    
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:left w:val="single" w:sz="24" w:space="0" w:color="319795"/><w:top w:val="none"/><w:right w:val="none"/><w:bottom w:val="none"/></w:tcBorders>') # Heavy teal left border
    tcPr.append(tcBorders)
    
    p_sum = cell.paragraphs[0]
    p_sum.paragraph_format.left_indent = Inches(0.15)
    p_sum.paragraph_format.right_indent = Inches(0.15)
    p_sum.paragraph_format.space_before = Pt(8)
    p_sum.paragraph_format.space_after = Pt(8)
    p_sum.paragraph_format.line_spacing = 1.15
    run_sum = p_sum.add_run(
        "As nonprofit organizations navigate a landscape of changing donor habits and stagnating traditional retention rates in 2026, "
        "Artificial Intelligence (AI) has emerged as a critical driver of operational capacity and strategic engagement. Rather than merely "
        "automating administrative workflows, leading NGOs are leveraging predictive AI and data-driven personalization to transition from generic mass appeals "
        "to deeper, individualized donor relations. This report examines current sector-wide adoption trends, analyzes key metrics, shares "
        "successful real-world case studies, and outlines responsible implementation strategies for forward-thinking NGOs like the InAmigos Foundation."
    )
    run_sum.font.size = Pt(10.5)
    run_sum.font.color.rgb = CHARCOAL
    run_sum.font.name = 'Calibri'
    run_sum.font.italic = True

    # Spacer
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # 1. AI Adoption Trends & Statistics in 2026
    add_custom_heading("1. AI Adoption Trends & Statistics", 1, 18, 6)
    
    add_body_paragraph(
        "The current year has marked a transition from experimental, ad-hoc AI usage to structured, CRM-integrated workflows. "
        "However, this widespread adoption highlights a unique dichotomy in how technology is utilized across the sector:"
    )
    
    add_bullet_point(
        " of nonprofits report using AI in their workflows in 2026. Despite this high figure, only 7% report that AI has fundamentally transformed their mission delivery or strategic outcomes. Most organizations remain on an 'efficiency plateau,' primarily employing AI for basic task acceleration (drafting emails, summarizing meeting logs, generating ideas).",
        bold_prefix="The Adoption Paradox (92% vs. 7%): 92%"
    )
    
    add_bullet_point(
        " of organizations report direct, measurable increases in fundraising revenue as a result of integrating AI tools. "
        "This indicates a significant return on investment for those who leverage AI beyond simple content drafting.",
        bold_prefix="Fundraising Revenue Impact: Approximately 30%"
    )
    
    add_bullet_point(
        "On average, nonprofit staff members utilizing integrated AI tools report saving between 15 to 20 hours per week "
        "on administrative data manipulation, audience-list creation, and report preparation. This time is directly reallocated "
        "toward high-value activities, such as major donor cultivation and community outreach.",
        bold_prefix="Operational Efficiency Gains: "
    )

    # 2. Donor Retention and Hyper-Personalization
    add_custom_heading("2. Optimizing Donor Retention & Engagement", 1, 18, 6)
    add_body_paragraph(
        "With recent donor participation numbers showing overall flat or downward trends globally, retaining active donors "
        "is a primary financial priority. Nonprofits are shifting focus from donor acquisition to donor retention using two AI methods:"
    )
    
    add_custom_heading("Predictive Churn Modeling", 2, 8, 4)
    add_body_paragraph(
        "By analyzing historical transaction records, engagement frequency, event attendance, and newsletter open rates, predictive AI algorithms "
        "calculate individual donor 'risk scores.' Fundraisers are automatically alerted when a previously active donor's behavior matches "
        "historical churn patterns (e.g., missed email openings, skipping a seasonal campaign). This allows staff to intervene with customized "
        "re-engagement strategies before the donor lapses completely."
    )
    
    add_custom_heading("Hyper-Personalization and Donor Trust", 2, 8, 4)
    add_body_paragraph(
        "Generic, one-size-fits-all fundraising appeals are increasingly ineffective. In 2026, AI-driven segmentation matches donor profiles "
        "with the exact projects they support (e.g., highlighting clean water outcomes to clean water advocates). However, trust is paramount:"
    )
    
    add_bullet_point(
        " of donors state that transparency regarding how an organization uses AI is highly important to them.",
        bold_prefix="Demanded Transparency: 93%"
    )
    add_bullet_point(
        " of donors express concern or discomfort regarding how their personal and financial data is analyzed by AI systems. "
        "This underscores the urgent need for clear data governance and privacy policies.",
        bold_prefix="Data Privacy Concerns: Nearly 40%"
    )

    # Page Break for a clean 2-page report
    doc.add_page_break()

    # 3. Real-world Case Studies
    add_custom_heading("3. Real-World Case Studies", 1, 12, 6)
    
    add_body_paragraph(
        "Real-world deployments demonstrate the potential of AI when paired with human intervention and empathetic storytelling."
    )
    
    add_custom_heading("Case Study 1: Greenpeace Australia Pacific — Predictive Outreach", 2, 8, 4)
    add_body_paragraph(
        "To combat donor attrition, Greenpeace Australia Pacific integrated a predictive machine learning tool (Dataro Predict) with their CRM. "
        "The model assigned real-time retention scores to active monthly donors, identifying those highly likely to churn. Instead of a broad, "
        "expensive direct-mail campaign, the organization conducted high-touch, personalized telephone calls exclusively to the at-risk segment. "
        "The results were stark:"
    )
    add_bullet_point("Targeted regular donors were 2.5 times less likely to cancel their subscriptions compared to control groups.")
    add_bullet_point("The targeted campaign generated $235,000 in saved donations over an 18-month tracking period.")
    add_bullet_point("The campaign achieved a return on investment (ROI) of over 10:1.")

    add_custom_heading("Case Study 2: Charity:Water — Immersive AI Storytelling", 2, 8, 4)
    add_body_paragraph(
        "Charity:Water deployed an interactive, conversational AI chatbot named 'Yeshi' via Facebook Messenger. Instead of asking for funds "
        "immediately, the AI walked users through the daily life of a woman in Ethiopia fetching water. The bot simulated real-time conversation, "
        "provided maps of the route, and shared photos and audio files. This storytelling approach transformed donor engagement:"
    )
    add_bullet_point("It built deep, cognitive empathy before asking for financial support.")
    add_bullet_point("It successfully engaged younger demographics who typically avoid email-based newsletters.")
    add_bullet_point("It established a conversational model for donor onboarding that continues to influence fundraising standards.")

    # 4. Comparative Metrics Table
    add_custom_heading("4. Key Performance Metrics for AI Adoption", 1, 18, 6)
    add_body_paragraph(
        "NGOs tracking their digital transformation should evaluate performance across these three main categories:"
    )

    # Table creation
    table = doc.add_table(rows=4, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    # Table header styling
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Category"
    hdr_cells[1].text = "Primary Metric"
    hdr_cells[2].text = "Target Benchmark (2026)"
    
    # Set widths
    col_widths = [Inches(1.8), Inches(2.2), Inches(2.5)]
    for row in table.rows:
        for idx, width in enumerate(col_widths):
            row.cells[idx].width = width
            
    # Set header colors & borders
    for cell in hdr_cells:
        tcPr = cell._tc.get_or_add_tcPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1A365D"/>') # Navy background
        tcPr.append(shd)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255) # White text
                run.font.size = Pt(10)
                run.font.name = 'Calibri'
                
    # Populate rows
    row_data = [
        ("Fundraising Effectiveness", "Donor Retention & Churn Reduction", "+15% increase in retention for targeted groups"),
        ("Operational Efficiency", "Administrative Time Saved", "15–20 hours saved weekly per staff member"),
        ("Strategic Impact", "Donor Engagement (Email Open/Click)", "Double the industry standard (>45% open rates)")
    ]
    
    for idx, data in enumerate(row_data):
        row = table.rows[idx + 1]
        row.cells[0].text = data[0]
        row.cells[1].text = data[1]
        row.cells[2].text = data[2]
        
        # Zebra striping for even rows
        bg_color = "F7FAFC" if idx % 2 == 0 else "FFFFFF"
        
        for c_idx, cell in enumerate(row.cells):
            tcPr = cell._tc.get_or_add_tcPr()
            shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg_color}"/>')
            tcPr.append(shd)
            
            # Subtly thin border gridlines
            tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:bottom w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/></w:tcBorders>')
            tcPr.append(tcBorders)
            
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)
                for run in p.runs:
                    run.font.size = Pt(9.5)
                    run.font.color.rgb = CHARCOAL
                    run.font.name = 'Calibri'

    p_spacer = doc.add_paragraph()
    p_spacer.paragraph_format.space_after = Pt(12)

    # 5. Strategic Recommendations for InAmigos Foundation
    add_custom_heading("5. Strategic Recommendations for InAmigos Foundation", 1, 12, 6)
    
    add_bullet_point(
        "Establish an ethical AI policy and share it publicly on your website. Since 93% of donors prize AI transparency, "
        "stating clearly how donor data is handled builds foundational trust.",
        bold_prefix="1. Build Transparency First: "
    )
    
    add_bullet_point(
        "Ensure CRM databases have clean, uniform, and deduplicated records. Predictive churn algorithms "
        "and personalization systems fail if underlying donor profiles contain errors or incomplete data.",
        bold_prefix="2. Audit CRM Data Quality: "
    )
    
    add_bullet_point(
        "Utilize AI-generated content (emails, social posts, letters) as initial drafts only. Incorporate "
        "human reviews for tone, local cultural relevance, and alignment with the foundation's core values before dispatch.",
        bold_prefix="3. Maintain the Human-in-the-Loop: "
    )
    
    add_bullet_point(
        "Introduce micro-experiments using free or low-cost AI tools (like ChatGPT for initial proposal brainstorming "
        "or small predictive scripts on donor spreadsheets) before investing in enterprise-grade platforms.",
        bold_prefix="4. Start with Focused Micro-Experiments: "
    )

    # Footer elements
    # Add page number or simple text footer
    for section in doc.sections:
        footer = section.footer
        p_foot = footer.paragraphs[0]
        p_foot.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_foot.paragraph_format.space_before = Pt(12)
        run_foot = p_foot.add_run("InAmigos Foundation | AI-Powered Data Analysis Report (2026)  —  Page ")
        run_foot.font.size = Pt(8.5)
        run_foot.font.color.rgb = SLATE
        run_foot.font.name = 'Calibri'
        
        # Word page field XML
        fldChar1 = parse_xml(r'<w:fldChar %s w:fldCharType="begin"/>' % nsdecls('w'))
        instrText = parse_xml(r'<w:instrText %s xml:space="preserve"> PAGE </w:instrText>' % nsdecls('w'))
        fldChar2 = parse_xml(r'<w:fldChar %s w:fldCharType="separate"/>' % nsdecls('w'))
        fldChar3 = parse_xml(r'<w:fldChar %s w:fldCharType="end"/>' % nsdecls('w'))
        p_foot._p.append(fldChar1)
        p_foot._p.append(instrText)
        p_foot._p.append(fldChar2)
        p_foot._p.append(fldChar3)

    # Save document
    filename = "AI_NGO_Fundraising_Report.docx"
    doc.save(filename)
    print(f"Report generated successfully: {os.path.abspath(filename)}")

if __name__ == "__main__":
    create_report()
